# -*- coding: utf-8 -*-
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

md_path = r'C:\Users\USER\music-streaming\diploma\thesis_complete.md'
with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code = False
code_buf = []
while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    if stripped.startswith('`'):
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_buf))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue
    if line.startswith('# ') and not line.startswith('## '):
        h = doc.add_heading(line[2:].strip(), level=1)
        for r in h.runs: r.font.name = 'Times New Roman'
    elif line.startswith('## '):
        h = doc.add_heading(line[3:].strip(), level=2)
        for r in h.runs: r.font.name = 'Times New Roman'
    elif line.startswith('### '):
        h = doc.add_heading(line[4:].strip(), level=3)
        for r in h.runs: r.font.name = 'Times New Roman'
    elif stripped == '---':
        pass
    elif stripped:
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
    i += 1

docx_path = r'C:\Users\USER\music-streaming\diploma\thesis_complete.docx'
doc.save(docx_path)
words = len(content.split())
print(f'DOCX: {words} words, ~{words//250} pages')
